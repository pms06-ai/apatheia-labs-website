"""
Apatheia Labs — Document Processing with Modal

This script runs on Modal's serverless infrastructure to process documents.
30 hours/month free tier is plenty for single-user usage.

Setup:
1. pip install modal
2. modal token new
3. modal deploy modal/process_pdf.py

Usage from Next.js:
- Call Modal's webhook endpoint with document data
- Or use Modal's Python client from a serverless function
"""

import modal

# Create Modal app
app = modal.App("apatheia-document-processor")

# Define the container image with required packages
image = modal.Image.debian_slim(python_version="3.11").pip_install(
    "pymupdf",           # PDF extraction
    "python-docx",       # DOCX parsing
    "marker-pdf",        # PDF to markdown (optional, heavy)
    "pillow",            # Image processing
    "httpx",             # HTTP client for callbacks
)


@app.function(
    image=image,
    timeout=300,
    memory=2048,
)
def extract_text_from_pdf(pdf_bytes: bytes) -> dict:
    """
    Extract text and metadata from a PDF file.
    
    Returns:
        {
            "text": str,
            "pages": int,
            "metadata": dict,
            "images": list[str],  # base64 encoded
        }
    """
    import fitz  # PyMuPDF
    import base64
    from io import BytesIO
    
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    pages_text = []
    images = []
    
    for page_num, page in enumerate(doc):
        # Extract text
        text = page.get_text("text")
        pages_text.append({
            "page": page_num + 1,
            "text": text,
        })
        
        # Extract images (optional, can be disabled for speed)
        for img_index, img in enumerate(page.get_images(full=True)):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_b64 = base64.b64encode(image_bytes).decode()
                images.append({
                    "page": page_num + 1,
                    "index": img_index,
                    "format": base_image["ext"],
                    "data": image_b64[:100] + "...",  # Truncated for response
                })
            except Exception:
                pass
    
    # Get metadata
    metadata = doc.metadata
    
    # Combine all text
    full_text = "\n\n".join([p["text"] for p in pages_text])
    
    doc.close()
    
    return {
        "text": full_text,
        "pages": len(pages_text),
        "pages_detail": pages_text,
        "metadata": metadata,
        "images_count": len(images),
    }


@app.function(
    image=image,
    timeout=120,
    memory=1024,
)
def extract_text_from_docx(docx_bytes: bytes) -> dict:
    """
    Extract text from a DOCX file.
    """
    from docx import Document
    from io import BytesIO
    
    doc = Document(BytesIO(docx_bytes))
    
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    full_text = "\n\n".join(paragraphs)
    
    return {
        "text": full_text,
        "paragraphs": len(paragraphs),
        "tables": len(tables),
        "tables_data": tables[:5],  # First 5 tables
    }


@app.function(
    image=modal.Image.debian_slim(python_version="3.11").pip_install(
        "faster-whisper",
        "torch",
    ),
    gpu="T4",  # Use GPU for transcription
    timeout=600,
    memory=8192,
)
def transcribe_audio(audio_bytes: bytes, language: str = "en") -> dict:
    """
    Transcribe audio using Faster-Whisper.
    
    Note: This uses GPU and may consume credits quickly.
    For free tier, consider using Replicate's Whisper instead.
    """
    from faster_whisper import WhisperModel
    from io import BytesIO
    import tempfile
    import os
    
    # Save to temp file (whisper needs file path)
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
        f.write(audio_bytes)
        temp_path = f.name
    
    try:
        # Load model
        model = WhisperModel("base", device="cuda", compute_type="float16")
        
        # Transcribe
        segments, info = model.transcribe(temp_path, language=language)
        
        # Collect segments
        transcript_segments = []
        full_text = []
        
        for segment in segments:
            transcript_segments.append({
                "start": segment.start,
                "end": segment.end,
                "text": segment.text,
            })
            full_text.append(segment.text)
        
        return {
            "text": " ".join(full_text),
            "language": info.language,
            "duration": info.duration,
            "segments": transcript_segments,
        }
    finally:
        os.unlink(temp_path)


@app.function(
    image=image,
    timeout=60,
)
def extract_entities_spacy(text: str) -> dict:
    """
    Extract named entities using spaCy.
    
    Note: For legal documents, consider using Blackstone model.
    """
    # This is a placeholder - you'd install spacy and a model
    # For production, use: pip_install("spacy", "en_core_web_sm")
    
    # Simple regex-based extraction as fallback
    import re
    
    entities = []
    
    # Dates
    date_pattern = r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b'
    for match in re.finditer(date_pattern, text):
        entities.append({"text": match.group(), "type": "DATE", "start": match.start()})
    
    # Potential names (capitalized words)
    name_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b'
    for match in re.finditer(name_pattern, text):
        entities.append({"text": match.group(), "type": "PERSON", "start": match.start()})
    
    return {
        "entities": entities[:100],  # Limit
        "count": len(entities),
    }


# Web endpoint for calling from Next.js
@app.function(image=image)
@modal.web_endpoint(method="POST")
def process_document(request: dict) -> dict:
    """
    Web endpoint to process a document.
    
    Request body:
    {
        "document_url": str,  # URL to fetch document from
        "document_type": str,  # "pdf" or "docx"
        "callback_url": str,  # Optional URL to POST results to
    }
    """
    import httpx
    
    doc_url = request.get("document_url")
    doc_type = request.get("document_type", "pdf")
    callback_url = request.get("callback_url")
    
    # Fetch document
    response = httpx.get(doc_url, timeout=60)
    doc_bytes = response.content
    
    # Process based on type
    if doc_type == "pdf":
        result = extract_text_from_pdf.remote(doc_bytes)
    elif doc_type == "docx":
        result = extract_text_from_docx.remote(doc_bytes)
    else:
        return {"error": f"Unsupported document type: {doc_type}"}
    
    # Send callback if provided
    if callback_url:
        httpx.post(callback_url, json=result, timeout=30)
    
    return result


if __name__ == "__main__":
    # Local testing
    with open("test.pdf", "rb") as f:
        result = extract_text_from_pdf.local(f.read())
        print(f"Extracted {result['pages']} pages")
        print(result['text'][:500])
